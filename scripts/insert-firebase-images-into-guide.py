from __future__ import annotations

from pathlib import Path
from typing import Iterable
import shutil

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(r"c:\TaskAm-main\TaskAm-main")
DOC_PATH = ROOT / "Intersite-Track-Complete-Guide.docx"
BACKUP_PATH = ROOT / "Intersite-Track-Complete-Guide.backup-before-firebase-images.docx"
ASSET_DIR = ROOT / "tmp" / "docs" / "firebase-guide-assets"

TAHOMA = Path(r"C:\Windows\Fonts\tahoma.ttf")
TAHOMA_BOLD = Path(r"C:\Windows\Fonts\tahomabd.ttf")

WIDTH = 1600
HEIGHT = 900
BG = "#F4F7FB"
PRIMARY = "#FFCA28"
PRIMARY_DARK = "#F9A825"
GREEN = "#2E7D32"
GREEN_LIGHT = "#E8F5E9"
BLUE = "#1976D2"
BLUE_LIGHT = "#E3F2FD"
PURPLE = "#6A1B9A"
PURPLE_LIGHT = "#F3E5F5"
RED = "#C62828"
RED_LIGHT = "#FFEBEE"
TEXT = "#1F2937"
MUTED = "#5B6470"
LINE = "#D0D7E2"
WHITE = "#FFFFFF"

FIGURES = [
    ("รูปที่ 1.1", "หน้าการสร้างโครงการใน Firebase Console"),
    ("รูปที่ 1.2", "หน้าดึง Service Account และค่า Web App Configuration"),
    ("รูปที่ 1.3", "หน้าการเปิดใช้งาน Email/Password Authentication"),
    ("รูปที่ 1.4", "หน้าการสร้าง Firestore Database และเตรียม Collection หลัก"),
    ("รูปที่ 1.5", "หน้าการกำหนด Security Rules และการตรวจสอบก่อนใช้งาน"),
]


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = TAHOMA_BOLD if bold else TAHOMA
    return ImageFont.truetype(str(path), size=size)


def rounded(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], fill: str, outline: str | None = None, radius: int = 24, width: int = 2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def text(draw: ImageDraw.ImageDraw, xy: tuple[int, int], value: str, size: int = 28, bold: bool = False, fill: str = TEXT):
    draw.text(xy, value, font=font(size, bold), fill=fill)


def multiline(draw: ImageDraw.ImageDraw, xy: tuple[int, int], lines: Iterable[str], size: int = 24, bold: bool = False, fill: str = TEXT, spacing: int = 14):
    draw.multiline_text(xy, "\n".join(lines), font=font(size, bold), fill=fill, spacing=spacing)


def browser_shell(draw: ImageDraw.ImageDraw, title: str):
    rounded(draw, (30, 30, WIDTH - 30, HEIGHT - 30), "#E6ECF5", radius=40)
    rounded(draw, (40, 40, WIDTH - 40, HEIGHT - 40), WHITE, outline="#BAC7D8", radius=36, width=3)
    rounded(draw, (40, 40, WIDTH - 40, 120), "#EEF3F8", radius=36)
    draw.rectangle((40, 80, WIDTH - 40, 120), fill="#EEF3F8")
    for i, color in enumerate(["#FF6B6B", "#FFD166", "#06D6A0"]):
        draw.ellipse((72 + i * 34, 72, 96 + i * 34, 96), fill=color)
    rounded(draw, (200, 64, WIDTH - 120, 98), WHITE, outline=LINE, radius=16)
    text(draw, (230, 68), title, size=22, fill=MUTED)


def sidebar(draw: ImageDraw.ImageDraw, items: list[tuple[str, bool]], x: int = 80, top: int = 160, width: int = 280):
    rounded(draw, (x, top, x + width, HEIGHT - 90), "#FAFBFD", outline=LINE, radius=28)
    y = top + 34
    text(draw, (x + 28, y), "Firebase", size=30, bold=True)
    y += 56
    for label, active in items:
        if active:
            rounded(draw, (x + 18, y - 8, x + width - 18, y + 36), PRIMARY, radius=18)
        text(draw, (x + 34, y), label, size=24, bold=active, fill=TEXT if active else MUTED)
        y += 58


def callout(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title_text: str, body_lines: list[str], fill: str = BLUE_LIGHT, title_fill: str = BLUE):
    rounded(draw, box, fill, outline=LINE, radius=28)
    x1, y1, _, _ = box
    text(draw, (x1 + 24, y1 + 20), title_text, size=28, bold=True, fill=title_fill)
    multiline(draw, (x1 + 24, y1 + 72), [f"- {line}" for line in body_lines], size=22, fill=TEXT, spacing=12)


def panel_title(draw: ImageDraw.ImageDraw, title_text: str, subtitle: str):
    text(draw, (370, 160), title_text, size=40, bold=True)
    text(draw, (370, 212), subtitle, size=24, fill=MUTED)


def save_step_1(path: Path):
    img = Image.new("RGB", (WIDTH, HEIGHT), BG)
    draw = ImageDraw.Draw(img)
    browser_shell(draw, "console.firebase.google.com")
    sidebar(draw, [("Build", False), ("Analytics", False), ("Project Overview", True)])
    panel_title(draw, "Step 1: Create Firebase Project", "สร้างโปรเจกต์ใหม่ใน Firebase Console")

    rounded(draw, (370, 280, 1180, 760), WHITE, outline=LINE, radius=30)
    text(draw, (410, 320), "Create project", size=32, bold=True)
    text(draw, (410, 376), "Project name", size=22, fill=MUTED)
    rounded(draw, (410, 412, 1140, 470), "#F9FBFD", outline=LINE, radius=16)
    text(draw, (436, 426), "internsite-track", size=24, bold=True)
    text(draw, (410, 510), "Google Analytics", size=22, fill=MUTED)
    rounded(draw, (410, 542, 1140, 636), GREEN_LIGHT, outline="#C8E6C9", radius=18)
    text(draw, (436, 566), "Disable Google Analytics for this project", size=24, bold=True, fill=GREEN)
    rounded(draw, (912, 680, 1140, 736), PRIMARY, outline=PRIMARY_DARK, radius=18)
    text(draw, (978, 694), "Create", size=26, bold=True)

    callout(draw, (1220, 300, 1510, 570), "ลำดับที่ควรทำ", [
        "เปิด Firebase Console",
        "กด Create a new project",
        "ตั้งชื่อโปรเจกต์",
        "ปิด Analytics ถ้าไม่จำเป็น",
        "ยืนยันการสร้างโปรเจกต์",
    ])

    img.save(path)


def save_step_2(path: Path):
    img = Image.new("RGB", (WIDTH, HEIGHT), BG)
    draw = ImageDraw.Draw(img)
    browser_shell(draw, "Project Settings")
    sidebar(draw, [("Project Overview", False), ("Project Settings", True), ("Service Accounts", True)])
    panel_title(draw, "Step 2: Get Credentials", "ดึง Service Account และ Web App Config")

    rounded(draw, (370, 280, 840, 760), WHITE, outline=LINE, radius=30)
    text(draw, (404, 316), "Service Accounts", size=30, bold=True)
    rounded(draw, (404, 378, 804, 446), "#F9FBFD", outline=LINE, radius=16)
    text(draw, (430, 400), "SDK: Node.js", size=24, bold=True)
    rounded(draw, (404, 490, 804, 572), PURPLE_LIGHT, outline="#E1BEE7", radius=18)
    text(draw, (430, 518), "Generate New Private Key", size=24, bold=True, fill=PURPLE)
    multiline(draw, (404, 610), [
        "ได้ไฟล์ JSON สำหรับ backend",
        "เก็บไฟล์ไว้นอก repo",
        "ใช้ค่า projectId / clientEmail / privateKey",
    ], size=22, fill=MUTED)

    rounded(draw, (872, 280, 1180, 760), WHITE, outline=LINE, radius=30)
    text(draw, (904, 316), "General > Your apps", size=28, bold=True)
    rounded(draw, (904, 378, 1148, 438), BLUE_LIGHT, outline="#BBDEFB", radius=16)
    text(draw, (934, 396), "Add Web app", size=24, bold=True, fill=BLUE)
    rounded(draw, (904, 476, 1148, 680), "#101828", radius=18)
    multiline(draw, (932, 504), [
        "apiKey: \"...\"",
        "authDomain: \"...\"",
        "projectId: \"...\"",
        "storageBucket: \"...\"",
        "messagingSenderId: \"...\"",
        "appId: \"...\"",
    ], size=21, fill="#E5E7EB", spacing=10)

    callout(draw, (1220, 300, 1510, 610), "ผลลัพธ์ที่ต้องคัดลอก", [
        "Service Account JSON",
        "firebaseConfig สำหรับ frontend",
        "ค่า projectId ให้ตรงกันทั้งระบบ",
        "ใช้ .env แยก backend กับ VITE_ frontend",
    ], fill=PURPLE_LIGHT, title_fill=PURPLE)

    img.save(path)


def save_step_3(path: Path):
    img = Image.new("RGB", (WIDTH, HEIGHT), BG)
    draw = ImageDraw.Draw(img)
    browser_shell(draw, "Authentication / Sign-in method")
    sidebar(draw, [("Authentication", True), ("Users", False), ("Templates", False)])
    panel_title(draw, "Step 3: Enable Authentication", "เปิด Email/Password สำหรับเข้าสู่ระบบ")

    rounded(draw, (370, 280, 1180, 760), WHITE, outline=LINE, radius=30)
    text(draw, (404, 316), "Sign-in providers", size=30, bold=True)

    providers = [
        ("Email/Password", True),
        ("Google", False),
        ("Phone", False),
        ("Anonymous", False),
    ]
    y = 396
    for name, enabled in providers:
        rounded(draw, (404, y, 1146, y + 82), "#FBFCFE", outline=LINE, radius=16)
        text(draw, (430, y + 24), name, size=24, bold=True)
        badge_fill = GREEN_LIGHT if enabled else RED_LIGHT
        badge_text = "Enabled" if enabled else "Disabled"
        badge_color = GREEN if enabled else RED
        rounded(draw, (928, y + 18, 1110, y + 62), badge_fill, outline=LINE, radius=16)
        text(draw, (968, y + 28), badge_text, size=22, bold=True, fill=badge_color)
        y += 96

    rounded(draw, (980, 680, 1146, 734), PRIMARY, outline=PRIMARY_DARK, radius=18)
    text(draw, (1044, 694), "Save", size=24, bold=True)

    callout(draw, (1220, 300, 1510, 560), "เช็กก่อนกด Save", [
        "เข้าเมนู Authentication",
        "เปิด Sign-in method",
        "Enable Email/Password",
        "กด Save",
    ], fill=GREEN_LIGHT, title_fill=GREEN)

    img.save(path)


def save_step_4(path: Path):
    img = Image.new("RGB", (WIDTH, HEIGHT), BG)
    draw = ImageDraw.Draw(img)
    browser_shell(draw, "Firestore Database")
    sidebar(draw, [("Firestore Database", True), ("Data", True), ("Rules", False)])
    panel_title(draw, "Step 4: Create Firestore Database", "สร้างฐานข้อมูลและเตรียม collection หลัก")

    rounded(draw, (370, 280, 780, 760), WHITE, outline=LINE, radius=30)
    text(draw, (404, 316), "Create database", size=30, bold=True)
    rounded(draw, (404, 390, 742, 470), BLUE_LIGHT, outline="#BBDEFB", radius=16)
    text(draw, (430, 418), "Start in test mode", size=24, bold=True, fill=BLUE)
    rounded(draw, (404, 504, 742, 584), GREEN_LIGHT, outline="#C8E6C9", radius=16)
    text(draw, (430, 532), "Region: asia-southeast1", size=24, bold=True, fill=GREEN)
    rounded(draw, (520, 664, 742, 720), PRIMARY, outline=PRIMARY_DARK, radius=18)
    text(draw, (590, 678), "Create", size=24, bold=True)

    rounded(draw, (812, 280, 1180, 760), WHITE, outline=LINE, radius=30)
    text(draw, (844, 316), "Collections", size=30, bold=True)
    for y, value in zip([396, 486, 576, 666], ["users", "tasks", "holidays", "app_settings"]):
        rounded(draw, (844, y, 1146, y + 64), "#F9FBFD", outline=LINE, radius=16)
        text(draw, (876, y + 18), value, size=24, bold=True)
    rounded(draw, (844, 616, 1146, 680), "#F9FBFD", outline=LINE, radius=16)
    text(draw, (876, 634), "saturday_schedules", size=24, bold=True)

    callout(draw, (1220, 300, 1510, 590), "ผลลัพธ์ที่ต้องได้", [
        "Firestore พร้อมใช้งาน",
        "เลือก region ใกล้ไทย",
        "มี collection หลักครบ",
        "ใช้ test mode ตอนเริ่มพัฒนา",
    ], fill=BLUE_LIGHT, title_fill=BLUE)

    img.save(path)


def save_step_5(path: Path):
    img = Image.new("RGB", (WIDTH, HEIGHT), BG)
    draw = ImageDraw.Draw(img)
    browser_shell(draw, "Firestore Rules / Common Checks")
    sidebar(draw, [("Firestore Database", False), ("Rules", True), ("Usage", False)])
    panel_title(draw, "Step 5: Security Rules & Validation", "ตั้ง Rules และเช็กปัญหาที่พบบ่อย")

    rounded(draw, (370, 280, 930, 760), WHITE, outline=LINE, radius=30)
    text(draw, (404, 316), "Rules Editor", size=30, bold=True)
    rounded(draw, (404, 378, 892, 700), "#111827", radius=18)
    multiline(draw, (434, 414), [
        "rules_version = '2';",
        "service cloud.firestore {",
        "  match /databases/{database}/documents {",
        "    match /users/{userId} {",
        "      allow read, write: if request.auth.uid == userId;",
        "    }",
        "    match /tasks/{document=**} {",
        "      allow read: if request.auth != null;",
        "      allow write: if request.auth.token.admin == true;",
        "    }",
        "  }",
        "}",
    ], size=20, fill="#E5E7EB", spacing=10)

    callout(draw, (972, 280, 1510, 760), "ก่อนทดสอบระบบ", [
        "เปิด Firestore API ใน Google Cloud",
        "ตรวจ VITE_FIREBASE_* ให้ครบ",
        "ใช้ default import ของ firebase-admin",
        "ทดสอบ login และ read/write rules",
    ], fill=RED_LIGHT, title_fill=RED)

    img.save(path)


def insert_paragraph_after(paragraph: Paragraph, text_value: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text_value:
        new_para.add_run(text_value)
    return new_para


def insert_page_break_before(paragraph: Paragraph):
    breaker = paragraph.insert_paragraph_before()
    breaker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    breaker.add_run().add_break(WD_BREAK.PAGE)
    return breaker


def style_heading(paragraph: Paragraph, size: int, bold: bool = True, color: tuple[int, int, int] = (31, 78, 121), center: bool = False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(10)
    for run in paragraph.runs:
        run.font.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(*color)


def style_normal(paragraph: Paragraph, size: int = 11, color: tuple[int, int, int] = (52, 64, 84)):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.left_indent = Pt(18)
    for run in paragraph.runs:
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(*color)


def apply_image_paragraph_style(paragraph: Paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "6")
        element.set(qn("w:color"), "D6DEE8")
        p_bdr.append(element)
    p_pr.append(p_bdr)


def add_image_block(after_paragraph: Paragraph, image_path: Path, caption: str):
    image_para = insert_paragraph_after(after_paragraph)
    apply_image_paragraph_style(image_para)
    image_run = image_para.add_run()
    image_run.add_picture(str(image_path), width=Inches(6.1))

    caption_para = insert_paragraph_after(image_para)
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.paragraph_format.space_before = Pt(4)
    caption_para.paragraph_format.space_after = Pt(16)
    run = caption_para.add_run(caption)
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(90, 98, 110)
    return caption_para


def find_paragraph(doc: Document, exact_text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def insert_figure_index(doc: Document):
    toc_anchor = find_paragraph(doc, "สารบัญ")
    current = insert_paragraph_after(toc_anchor)
    current.add_run("สารบัญรูปภาพ")
    style_heading(current, size=14, color=(31, 78, 121))

    for number, title in FIGURES:
        current = insert_paragraph_after(current)
        current.add_run(f"{number} {title}")
        style_normal(current, size=10.5)

    current = insert_paragraph_after(current)
    current.paragraph_format.space_after = Pt(10)
    return current


def style_section_layout(doc: Document):
    for heading in ["PART 1: Firebase Setup & Configuration", "PART 2: Intersite Track — User Guide"]:
        insert_page_break_before(find_paragraph(doc, heading))
        para = find_paragraph(doc, heading)
        style_heading(para, size=18, center=True)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(14)

    for heading in ["1.1 Firebase Console Setup", "1.2 Authentication Setup", "1.3 Firestore Database Setup", "1.4 Security Rules", "1.5 Common Firebase Issues & Fixes"]:
        insert_page_break_before(find_paragraph(doc, heading))
        para = find_paragraph(doc, heading)
        style_heading(para, size=15)
        para.paragraph_format.space_after = Pt(12)

    for heading in ["2.1 ภาพรวมระบบ", "2.2 LINE Bot Setup"]:
        para = find_paragraph(doc, heading)
        style_heading(para, size=15)
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after = Pt(10)


def main():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    assets = {
        "create_project": ASSET_DIR / "firebase-step-1-create-project.png",
        "credentials": ASSET_DIR / "firebase-step-2-credentials.png",
        "authentication": ASSET_DIR / "firebase-step-3-authentication.png",
        "firestore": ASSET_DIR / "firebase-step-4-firestore.png",
        "rules": ASSET_DIR / "firebase-step-5-rules.png",
    }

    save_step_1(assets["create_project"])
    save_step_2(assets["credentials"])
    save_step_3(assets["authentication"])
    save_step_4(assets["firestore"])
    save_step_5(assets["rules"])

    if not BACKUP_PATH.exists():
      shutil.copy2(DOC_PATH, BACKUP_PATH)

    doc = Document(str(BACKUP_PATH if BACKUP_PATH.exists() else DOC_PATH))

    insert_figure_index(doc)
    style_section_layout(doc)

    anchor1 = find_paragraph(doc, "กด Create project → รอ 1-2 นาที")
    add_image_block(anchor1, assets["create_project"], f"{FIGURES[0][0]} {FIGURES[0][1]}")

    anchor2 = find_paragraph(doc, "Copy config object (apiKey, authDomain, projectId, etc.)")
    add_image_block(anchor2, assets["credentials"], f"{FIGURES[1][0]} {FIGURES[1][1]}")

    anchor3 = find_paragraph(doc, "กด Email/Password → Enable → Save")
    add_image_block(anchor3, assets["authentication"], f"{FIGURES[2][0]} {FIGURES[2][1]}")

    anchor4 = find_paragraph(doc, "กด Create")
    # choose first occurrence after Firestore heading
    firestore_anchor = None
    passed_firestore_heading = False
    for paragraph in doc.paragraphs:
        value = paragraph.text.strip()
        if value == "1.3 Firestore Database Setup":
            passed_firestore_heading = True
        if passed_firestore_heading and value == "กด Create":
            firestore_anchor = paragraph
            break
    if firestore_anchor is None:
        firestore_anchor = anchor4
    add_image_block(firestore_anchor, assets["firestore"], f"{FIGURES[3][0]} {FIGURES[3][1]}")

    anchor5 = find_paragraph(doc, "ตั้ง test mode หรือ customize rules ให้อนุญาต")
    add_image_block(anchor5, assets["rules"], f"{FIGURES[4][0]} {FIGURES[4][1]}")

    doc.save(str(DOC_PATH))
    print(f"Updated: {DOC_PATH}")
    print(f"Backup: {BACKUP_PATH}")
    for key, path in assets.items():
        print(f"Asset {key}: {path}")


if __name__ == "__main__":
    main()
