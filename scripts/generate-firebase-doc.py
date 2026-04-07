"""
generate-firebase-doc.py
สร้าง Firebase_Tutorial.docx — คู่มือการตั้งค่าและใช้งาน Firebase กับโปรเจค Intersite Track

Usage:
    pip install python-docx
    python scripts/generate-firebase-doc.py
"""

import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_FILE = "Firebase_Tutorial.docx"

doc = Document()

# ─── Styles ──────────────────────────────────────────────────────────────────

def set_heading(paragraph, text, level=1):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(text)
    if not paragraph.runs:
        run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB) if level == 1 else RGBColor(0x1E, 0x40, 0xAF)
    run.font.size = Pt(18 if level == 1 else 14)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18 if level == 1 else (14 if level == 2 else 12))
    if level == 1:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    else:
        run.font.color.rgb = RGBColor(0x37, 0x51, 0x8C)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_step(doc, step_num, text):
    p = doc.add_paragraph()
    run = p.add_run(f"ขั้นตอนที่ {step_num}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
    p.add_run(text)
    return p


def add_code(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    # Light gray background via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F3F4F6")
    pPr.append(shd)
    return p


def add_screenshot_placeholder(doc, caption=""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f"📸  [ วางภาพ screenshot ที่นี่{' — ' + caption if caption else ''} ]")
    run.bold = True
    run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
    run.font.size = Pt(11)
    return p


def add_divider(doc):
    doc.add_paragraph("─" * 60)


# ─── Cover page ──────────────────────────────────────────────────────────────

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("คู่มือการตั้งค่าและใช้งาน Firebase")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_p.add_run("โปรเจค Intersite Track — CWIE Internship").font.size = Pt(14)

doc.add_paragraph()
info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_p.add_run("Stack: React 19 + Express + Firebase (Auth + Firestore)").font.size = Pt(11)

doc.add_page_break()

# ─── Section 1: Firebase คืออะไร ─────────────────────────────────────────────

add_heading(doc, "1. Firebase คืออะไร", level=1)
add_body(doc,
    "Firebase คือ Backend-as-a-Service (BaaS) จาก Google ที่ให้บริการโครงสร้างพื้นฐานสำหรับแอปพลิเคชัน "
    "โดยในโปรเจคนี้เราใช้ Firebase 3 บริการหลัก:"
)
add_body(doc, "• Firebase Authentication — จัดการระบบ Login/Logout ด้วย Email & Password")
add_body(doc, "• Cloud Firestore — ฐานข้อมูล NoSQL สำหรับเก็บข้อมูล tasks, users, notifications ฯลฯ")
add_body(doc, "• Firebase Admin SDK — ใช้ฝั่ง Backend (Node.js/Express) สำหรับ verify token และจัดการ users")
doc.add_paragraph()

# ─── Section 2: สร้าง Firebase Project ──────────────────────────────────────

add_heading(doc, "2. การสร้าง Firebase Project", level=1)
add_body(doc, "ทำตามขั้นตอนต่อไปนี้เพื่อสร้าง Firebase Project ใหม่:")

add_step(doc, 1, "เปิด Browser ไปที่ https://console.firebase.google.com")
add_step(doc, 2, "คลิกปุ่ม \"Add project\" หรือ \"Create a project\"")
add_screenshot_placeholder(doc, "Firebase Console — หน้าแรก")

add_step(doc, 3, "ตั้งชื่อ Project ว่า \"intersite-track\" แล้วคลิก Continue")
add_screenshot_placeholder(doc, "ตั้งชื่อ Project")

add_step(doc, 4, "เลือก Disable Google Analytics (ไม่จำเป็นสำหรับโปรเจคนี้) แล้วคลิก Create project")
add_screenshot_placeholder(doc, "Disable Analytics")

add_step(doc, 5, "รอ Firebase สร้าง Project เสร็จ แล้วคลิก Continue")
add_screenshot_placeholder(doc, "Project พร้อมใช้งาน")
doc.add_paragraph()

# ─── Section 3: ตั้งค่า Firebase Authentication ──────────────────────────────

add_heading(doc, "3. ตั้งค่า Firebase Authentication", level=1)
add_body(doc, "เปิดใช้งาน Email/Password Authentication:")

add_step(doc, 1, "ใน Firebase Console คลิกเมนู \"Authentication\" ทางซ้าย")
add_step(doc, 2, "คลิกแท็บ \"Sign-in method\"")
add_screenshot_placeholder(doc, "Authentication → Sign-in method")

add_step(doc, 3, "คลิก \"Email/Password\" แล้วเปิด Toggle ให้เป็น Enabled")
add_step(doc, 4, "คลิก Save")
add_screenshot_placeholder(doc, "เปิด Email/Password provider")

add_body(doc, "✅ Authentication พร้อมใช้งานแล้ว")
doc.add_paragraph()

# ─── Section 4: ตั้งค่า Firestore Database ───────────────────────────────────

add_heading(doc, "4. ตั้งค่า Firestore Database", level=1)
add_body(doc, "สร้าง Firestore Database สำหรับเก็บข้อมูลทั้งหมดของโปรเจค:")

add_step(doc, 1, "คลิกเมนู \"Firestore Database\" ทางซ้าย")
add_step(doc, 2, "คลิกปุ่ม \"Create database\"")
add_screenshot_placeholder(doc, "Firestore — Create database")

add_step(doc, 3, "เลือก \"Start in test mode\" (สำหรับ Development) แล้วคลิก Next")
add_screenshot_placeholder(doc, "Test mode")

add_step(doc, 4, "เลือก Location: \"asia-southeast1 (Singapore)\" แล้วคลิก Enable")
add_screenshot_placeholder(doc, "เลือก Region")

add_body(doc, "✅ Firestore Database พร้อมใช้งานแล้ว")
add_body(doc, "⚠️ หมายเหตุ: สำหรับ Production ให้เปลี่ยน Security Rules เป็น:")
add_code(doc,
    "rules_version = '2';\n"
    "service cloud.firestore {\n"
    "  match /databases/{database}/documents {\n"
    "    match /{document=**} {\n"
    "      allow read, write: if request.auth != null;\n"
    "    }\n"
    "  }\n"
    "}"
)
doc.add_paragraph()

# ─── Section 5: ดาวน์โหลด Service Account Key ────────────────────────────────

add_heading(doc, "5. ดาวน์โหลด Service Account Key (สำหรับ Backend)", level=1)
add_body(doc, "Backend (Express) ต้องการ Service Account Key เพื่อใช้ Firebase Admin SDK:")

add_step(doc, 1, "คลิกไอคอนฟันเฟือง ⚙️ บน Project Overview แล้วเลือก \"Project settings\"")
add_screenshot_placeholder(doc, "Project Settings")

add_step(doc, 2, "คลิกแท็บ \"Service accounts\"")
add_step(doc, 3, "คลิก \"Generate new private key\" แล้วยืนยัน")
add_screenshot_placeholder(doc, "Generate private key")

add_step(doc, 4, "ไฟล์ JSON จะถูกดาวน์โหลด — เก็บไว้อย่างปลอดภัย อย่า commit ขึ้น Git!")
add_body(doc, "จากนั้นเปิดไฟล์ JSON และคัดลอกค่าต่อไปนี้ใส่ใน .env ของโปรเจค:")
add_code(doc,
    "FIREBASE_PROJECT_ID=your-project-id\n"
    "FIREBASE_CLIENT_EMAIL=firebase-adminsdk-xxxxx@your-project-id.iam.gserviceaccount.com\n"
    "FIREBASE_PRIVATE_KEY=\"-----BEGIN PRIVATE KEY-----\\nyour-private-key\\n-----END PRIVATE KEY-----\\n\""
)
doc.add_paragraph()

# ─── Section 6: ดาวน์โหลด Web App Config (สำหรับ Frontend) ──────────────────

add_heading(doc, "6. ตั้งค่า Web App สำหรับ Frontend", level=1)
add_body(doc, "Frontend (React) ต้องการ Firebase Config object:")

add_step(doc, 1, "ใน Project Settings คลิกแท็บ \"General\"")
add_step(doc, 2, "เลื่อนลงไปที่ \"Your apps\" แล้วคลิกไอคอน </> เพื่อ Add Web App")
add_screenshot_placeholder(doc, "Add Web App")

add_step(doc, 3, "ตั้งชื่อ App ว่า \"intersite-track-web\" แล้วคลิก Register app")
add_step(doc, 4, "คัดลอก firebaseConfig object ที่ปรากฏ:")
add_screenshot_placeholder(doc, "Firebase Config object")

add_body(doc, "เพิ่มค่าเหล่านี้ใน .env:")
add_code(doc,
    "VITE_FIREBASE_API_KEY=AIzaSy...\n"
    "VITE_FIREBASE_AUTH_DOMAIN=intersite-track.firebaseapp.com\n"
    "VITE_FIREBASE_PROJECT_ID=intersite-track\n"
    "VITE_FIREBASE_STORAGE_BUCKET=intersite-track.firebasestorage.app\n"
    "VITE_FIREBASE_MESSAGING_SENDER_ID=123456789\n"
    "VITE_FIREBASE_APP_ID=1:123456789:web:abc123"
)
doc.add_paragraph()

# ─── Section 7: ติดตั้ง Firebase SDK ─────────────────────────────────────────

add_heading(doc, "7. ติดตั้ง Firebase SDK", level=1)
add_body(doc, "รันคำสั่งต่อไปนี้ใน Terminal ที่ root ของโปรเจค:")

add_code(doc, "npm install firebase firebase-admin")

add_body(doc, "• firebase — สำหรับ Frontend (React): Auth, Firestore client")
add_body(doc, "• firebase-admin — สำหรับ Backend (Express): Admin SDK, Firestore server-side")
doc.add_paragraph()

# ─── Section 8: การเชื่อมต่อ Frontend ───────────────────────────────────────

add_heading(doc, "8. การเชื่อมต่อ Frontend (React)", level=1)
add_body(doc, "ไฟล์: src/lib/firebase.ts — initialize Firebase App สำหรับ Frontend:")

add_code(doc,
    "import { initializeApp } from \"firebase/app\";\n"
    "import { getAuth } from \"firebase/auth\";\n"
    "import { getFirestore } from \"firebase/firestore\";\n"
    "\n"
    "const firebaseConfig = {\n"
    "  apiKey: import.meta.env.VITE_FIREBASE_API_KEY,\n"
    "  authDomain: import.meta.env.VITE_FIREBASE_AUTH_DOMAIN,\n"
    "  projectId: import.meta.env.VITE_FIREBASE_PROJECT_ID,\n"
    "  storageBucket: import.meta.env.VITE_FIREBASE_STORAGE_BUCKET,\n"
    "  messagingSenderId: import.meta.env.VITE_FIREBASE_MESSAGING_SENDER_ID,\n"
    "  appId: import.meta.env.VITE_FIREBASE_APP_ID,\n"
    "};\n"
    "\n"
    "export const firebaseApp = initializeApp(firebaseConfig);\n"
    "export const auth = getAuth(firebaseApp);\n"
    "export const firestore = getFirestore(firebaseApp);"
)

add_body(doc, "การใช้งาน Firebase Auth ใน Frontend (src/services/authService.ts):")
add_code(doc,
    "import { signInWithEmailAndPassword, signOut } from \"firebase/auth\";\n"
    "import { auth } from \"../lib/firebase\";\n"
    "\n"
    "// Login\n"
    "const credential = await signInWithEmailAndPassword(auth, email, password);\n"
    "const token = await credential.user.getIdToken(); // ใช้ token นี้ส่งไป Backend\n"
    "\n"
    "// Logout\n"
    "await signOut(auth);\n"
    "\n"
    "// Get current user token (สำหรับ API calls)\n"
    "const token = await auth.currentUser?.getIdToken();"
)
doc.add_paragraph()

# ─── Section 9: การเชื่อมต่อ Backend ────────────────────────────────────────

add_heading(doc, "9. การเชื่อมต่อ Backend (Firebase Admin SDK)", level=1)
add_body(doc, "ไฟล์: server/config/firebase-admin.ts — initialize Firebase Admin สำหรับ Backend:")

add_code(doc,
    "import * as admin from \"firebase-admin\";\n"
    "import dotenv from \"dotenv\";\n"
    "\n"
    "dotenv.config();\n"
    "\n"
    "if (!admin.apps.length) {\n"
    "  admin.initializeApp({\n"
    "    credential: admin.credential.cert({\n"
    "      projectId: process.env.FIREBASE_PROJECT_ID,\n"
    "      clientEmail: process.env.FIREBASE_CLIENT_EMAIL,\n"
    "      privateKey: process.env.FIREBASE_PRIVATE_KEY?.replace(/\\\\n/g, \"\\n\"),\n"
    "    }),\n"
    "  });\n"
    "}\n"
    "\n"
    "export const adminAuth = admin.auth();\n"
    "export const db = admin.firestore();\n"
    "export const FieldValue = admin.firestore.FieldValue;"
)

add_body(doc, "Auth Middleware — ตรวจสอบ Firebase Token ที่ส่งมาจาก Frontend:")
add_code(doc,
    "import { adminAuth, db } from \"../config/firebase-admin.js\";\n"
    "\n"
    "export async function requireAuth(req, res, next) {\n"
    "  const token = req.headers.authorization?.substring(7);\n"
    "  if (!token) { res.status(401).json({ error: \"กรุณาเข้าสู่ระบบก่อน\" }); return; }\n"
    "\n"
    "  const decodedToken = await adminAuth.verifyIdToken(token);\n"
    "  const uid = decodedToken.uid;\n"
    "\n"
    "  const userDoc = await db.collection(\"users\").doc(uid).get();\n"
    "  req.user = { id: uid, ...userDoc.data() };\n"
    "  next();\n"
    "}"
)
doc.add_paragraph()

# ─── Section 10: ตัวอย่างการใช้ Firestore ────────────────────────────────────

add_heading(doc, "10. ตัวอย่างการใช้ Firestore (Backend)", level=1)
add_body(doc, "การอ่านข้อมูล:")

add_code(doc,
    "import { db } from \"../config/firebase-admin.js\";\n"
    "\n"
    "// อ่าน document เดียว\n"
    "const doc = await db.collection(\"tasks\").doc(taskId).get();\n"
    "const task = doc.exists ? { id: doc.id, ...doc.data() } : null;\n"
    "\n"
    "// อ่านทั้ง collection\n"
    "const snap = await db.collection(\"tasks\").orderBy(\"created_at\", \"desc\").get();\n"
    "const tasks = snap.docs.map(doc => ({ id: doc.id, ...doc.data() }));\n"
    "\n"
    "// Query — หา tasks ของ user คนนึง\n"
    "const snap = await db.collection(\"tasks\")\n"
    "  .where(\"assignees\", \"array-contains\", userId)\n"
    "  .get();"
)

add_body(doc, "การเขียน/อัพเดทข้อมูล:")
add_code(doc,
    "// สร้าง document ใหม่ (auto-generated ID)\n"
    "const ref = db.collection(\"tasks\").doc();\n"
    "await ref.set({ title: \"งานใหม่\", status: \"pending\", created_at: new Date().toISOString() });\n"
    "const newTaskId = ref.id;\n"
    "\n"
    "// อัพเดท field บางส่วน\n"
    "await db.collection(\"tasks\").doc(taskId).update({ status: \"completed\" });\n"
    "\n"
    "// ลบ document\n"
    "await db.collection(\"tasks\").doc(taskId).delete();\n"
    "\n"
    "// Batch write (หลาย operations พร้อมกัน)\n"
    "const batch = db.batch();\n"
    "batch.update(db.collection(\"tasks\").doc(id1), { status: \"completed\" });\n"
    "batch.update(db.collection(\"tasks\").doc(id2), { status: \"in_progress\" });\n"
    "await batch.commit();"
)
doc.add_paragraph()

# ─── Section 11: ตัวอย่างการใช้ Firebase Auth ────────────────────────────────

add_heading(doc, "11. ตัวอย่างการใช้ Firebase Auth", level=1)

add_heading(doc, "Backend — สร้างและจัดการ Users", level=2)
add_code(doc,
    "import { adminAuth } from \"../config/firebase-admin.js\";\n"
    "\n"
    "// สร้าง user ใหม่\n"
    "const userRecord = await adminAuth.createUser({\n"
    "  email: \"user@example.com\",\n"
    "  password: \"password123\",\n"
    "  emailVerified: true,\n"
    "});\n"
    "const uid = userRecord.uid;\n"
    "\n"
    "// ลบ user\n"
    "await adminAuth.deleteUser(uid);\n"
    "\n"
    "// เปลี่ยน password\n"
    "await adminAuth.updateUser(uid, { password: \"newPassword\" });\n"
    "\n"
    "// Verify ID Token จาก Frontend\n"
    "const decoded = await adminAuth.verifyIdToken(idToken);\n"
    "const userId = decoded.uid;"
)

add_heading(doc, "Frontend — Login / Logout / Token", level=2)
add_code(doc,
    "import { signInWithEmailAndPassword, signOut, onAuthStateChanged } from \"firebase/auth\";\n"
    "import { auth } from \"../lib/firebase\";\n"
    "\n"
    "// Login\n"
    "const cred = await signInWithEmailAndPassword(auth, email, password);\n"
    "const idToken = await cred.user.getIdToken(); // ส่งเป็น Bearer token ไป Backend\n"
    "\n"
    "// Logout\n"
    "await signOut(auth);\n"
    "\n"
    "// ติดตาม Auth State\n"
    "onAuthStateChanged(auth, (user) => {\n"
    "  if (user) {\n"
    "    console.log(\"Logged in:\", user.uid);\n"
    "  } else {\n"
    "    console.log(\"Logged out\");\n"
    "  }\n"
    "});"
)

add_body(doc, "การส่ง ID Token ไปกับ API request:")
add_code(doc,
    "const token = await auth.currentUser?.getIdToken();\n"
    "const response = await fetch(\"/api/tasks\", {\n"
    "  headers: { Authorization: `Bearer ${token}` },\n"
    "});"
)
doc.add_paragraph()

# ─── Section 12: ขั้นตอน Checklist ──────────────────────────────────────────

add_heading(doc, "12. Checklist การตั้งค่า Firebase ทั้งหมด", level=1)

steps = [
    "สร้าง Firebase Project ใน console.firebase.google.com",
    "เปิด Authentication → Sign-in method → Email/Password",
    "เปิด Firestore Database → test mode → region: asia-southeast1",
    "ดาวน์โหลด Service Account Key จาก Project Settings → Service accounts",
    "เพิ่ม Web App และ copy firebaseConfig",
    "เพิ่มค่า ENV ทั้งหมดใน .env (FIREBASE_PROJECT_ID, FIREBASE_CLIENT_EMAIL, FIREBASE_PRIVATE_KEY, VITE_FIREBASE_*)",
    "รัน npm install firebase firebase-admin",
    "ตรวจสอบว่า npm run dev ทำงานได้ปกติ",
    "ทดสอบ Login/Logout ผ่านหน้าเว็บ",
    "ตรวจสอบใน Firestore Console ว่ามีข้อมูล users เพิ่มขึ้นหลัง Login",
    "เพิ่ม LINE_GROUP_ID ใน .env และทดสอบการแจ้งเตือน LINE Group",
]

for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"☐  {i}. {step}")

doc.add_paragraph()
add_divider(doc)
add_body(doc, "จัดทำโดย: ทีม Intersite Track | CWIE Internship 2026")

# ─── Section: LINE Messaging API ─────────────────────────────────────────────
add_heading(doc, "7. การตั้งค่า LINE Messaging API", level=1)
add_body(doc, "ระบบ Intersite Track ใช้ LINE Messaging API เพื่อส่งการแจ้งเตือนไปยังผู้ใช้งานผ่าน LINE")

add_heading(doc, "7.1 ขั้นตอนสร้าง LINE Channel", level=2)
steps_line = [
    "เข้า https://developers.line.biz/ และ Login ด้วย LINE account",
    "เลือก Provider (organization) ที่ต้องการ",
    "กด 'Create a Messaging API channel'",
    "กรอกข้อมูล Channel name, description, category",
    "กด 'Create' เพื่อสร้าง Channel",
]
for i, step in enumerate(steps_line, 1):
    add_body(doc, f"{i}. {step}")

add_heading(doc, "7.2 การได้ Channel Access Token", level=2)
add_body(doc, "ไปที่ Channel Settings → Messaging API tab → Channel access token → กด 'Issue'")
add_code(doc, "LINE_CHANNEL_ACCESS_TOKEN=your-token-here")

add_heading(doc, "7.3 การรับ LINE User ID ของ Admin", level=2)
add_body(doc, "เพิ่ม bot เป็นเพื่อนใน LINE แล้วส่งข้อความอะไรก็ได้ จากนั้นดู server logs จะเห็น userId")
add_code(doc, "LINE_ADMIN_USER_ID=Uxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx")

add_heading(doc, "7.4 การรับ LINE Group ID อัตโนมัติ", level=2)
add_body(doc, "ระบบรับ Group ID อัตโนมัติผ่าน LINE Webhook:")
group_steps = [
    "เปิด 'Allow bot to join group chats' ใน LINE Official Account Manager → Settings → Messaging API",
    "ตั้งค่า Webhook URL: https://your-domain.com/api/line/webhook",
    "สร้าง LINE Group และเพิ่ม bot เข้า group",
    "เมื่อ bot เข้า group ระบบจะบันทึก Group ID ลง Firestore อัตโนมัติ",
]
for i, step in enumerate(group_steps, 1):
    add_body(doc, f"{i}. {step}")

add_heading(doc, "7.5 การทดสอบ LINE API", level=2)
add_code(doc, 'curl -X POST https://api.line.me/v2/bot/message/push \\\n  -H "Authorization: Bearer YOUR_TOKEN" \\\n  -H "Content-Type: application/json" \\\n  -d \'{"to":"USER_ID","messages":[{"type":"text","text":"Test"}]}\'')

# ─── Section: วันหยุดและเวรเสาร์ ─────────────────────────────────────────────
add_heading(doc, "8. ระบบวันหยุดและตารางเวรเสาร์", level=1)

add_heading(doc, "8.1 วันหยุดประจำปี (Holidays)", level=2)
add_body(doc, "Firestore collection: holidays")
add_code(doc, '{\n  id: string,\n  date: "YYYY-MM-DD",\n  name: string,\n  type: "holiday" | "special",\n  created_at: string,\n  created_by: string\n}')
add_body(doc, "API Endpoints:")
add_code(doc, "GET  /api/holidays?year=2026&month=4\nPOST /api/holidays  (admin only)\nPUT  /api/holidays/:id  (admin only)\nDEL  /api/holidays/:id  (admin only)")

add_heading(doc, "8.2 ตารางเวรเสาร์ (Saturday Schedule)", level=2)
add_body(doc, "Firestore collection: saturday_schedules")
add_code(doc, '{\n  id: string,\n  date: "YYYY-MM-DD",  // must be Saturday\n  user_ids: string[],\n  note: string | null,\n  created_at: string,\n  created_by: string\n}')

add_heading(doc, "8.3 การแจ้งเตือน LINE อัตโนมัติ", level=2)
schedule_table_data = [
    ("ทุกวัน 08:00", "แจ้งวันหยุดวันนี้ + deadline งาน"),
    ("ทุกวัน 20:00", "แจ้งวันหยุดพรุ่งนี้"),
    ("ทุกวันจันทร์ 08:00", "สรุปวันหยุดสัปดาห์นี้"),
    ("ทุกวันศุกร์ 18:00", "แจ้งเวรเสาร์พรุ่งนี้ (เฉพาะผู้มีเวร)"),
    ("ทุกวันเสาร์ 08:00", "แจ้งเวรเสาร์วันนี้ (ส่วนตัว + group)"),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "เวลา"
hdr[1].text = "การแจ้งเตือน"
for time_str, desc in schedule_table_data:
    row = table.add_row().cells
    row[0].text = time_str
    row[1].text = desc
doc.add_paragraph()

# ─── Section: Environment Variables ──────────────────────────────────────────
add_heading(doc, "9. Environment Variables (.env)", level=1)
add_code(doc, """# Firebase Admin SDK
FIREBASE_PROJECT_ID=your-project-id
FIREBASE_CLIENT_EMAIL=firebase-adminsdk-xxx@project.iam.gserviceaccount.com
FIREBASE_PRIVATE_KEY="-----BEGIN PRIVATE KEY-----\\nyour-key\\n-----END PRIVATE KEY-----\\n"

# Firebase JS SDK (Frontend)
VITE_FIREBASE_API_KEY=your-api-key
VITE_FIREBASE_AUTH_DOMAIN=project.firebaseapp.com
VITE_FIREBASE_PROJECT_ID=your-project-id
VITE_FIREBASE_STORAGE_BUCKET=project.firebasestorage.app
VITE_FIREBASE_MESSAGING_SENDER_ID=your-sender-id
VITE_FIREBASE_APP_ID=1:sender-id:web:app-id

# LINE Messaging API
LINE_CHANNEL_ACCESS_TOKEN=your-token
LINE_ADMIN_USER_ID=Uxxxxxxxxxxxxxxx
LINE_GROUP_ID=  # ไม่ต้องใส่ — ระบบรับจาก webhook อัตโนมัติ

# Application
NODE_ENV=development
PORT=3694""")

# ─── Save ─────────────────────────────────────────────────────────────────────

doc.save(OUTPUT_FILE)
print(f"✅ สร้างไฟล์ {OUTPUT_FILE} เรียบร้อยแล้ว")
print(f"   ขนาดไฟล์: {os.path.getsize(OUTPUT_FILE):,} bytes")
